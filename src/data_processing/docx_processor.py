"""
DOCX processor — parallel clean + redline processing.

In M&A, negotiation history (redlines) is often the most diligence-relevant
content — struck indemnification caps, added representations, renegotiated
earnouts. Silently collapsing to final version discards critical findings.

Produces two parallel chunk sets:
- Clean: tracked changes accepted, is_redline=0
- Redline: additions (+) and deletions (~~strikethrough~~) inline, is_redline=1
"""

from pathlib import Path
from dataclasses import dataclass
from lxml import etree

import docx

from src.utils.logger import setup_logger

logger = setup_logger(__name__)

# Word XML namespaces
WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NSMAP = {"w": WORD_NS}


@dataclass
class Chunk:
    """A text chunk from a DOCX document."""
    text: str
    section_heading: str = ""
    page_number: int = 0
    is_redline: int = 0
    redline_base_doc_id: str = ""
    content_type: str = "text"
    metadata: dict = None

    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


def _extract_clean_text(doc: docx.Document) -> list[Chunk]:
    """
    Extracts clean text with tracked changes accepted.
    Insertions are included, deletions are excluded.

    Args:
        doc: Opened python-docx Document.

    Returns:
        List of Chunks with clean text and section headings.
    """
    chunks = []
    current_heading = ""

    for para in doc.paragraphs:
        # Detect headings
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            current_heading = para.text.strip()
            continue

        text = para.text.strip()
        if not text:
            continue

        chunks.append(Chunk(
            text=text,
            section_heading=current_heading,
            is_redline=0,
            content_type="text",
        ))

    return chunks


def _extract_redline_text(doc: docx.Document) -> list[Chunk]:
    """
    Extracts text with tracked changes shown inline:
    - Insertions marked with (+added text)
    - Deletions marked with (~~deleted text~~)

    Args:
        doc: Opened python-docx Document.

    Returns:
        List of Chunks with redline markup.
    """
    chunks = []
    current_heading = ""

    for para in doc.paragraphs:
        # Detect headings
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            current_heading = para.text.strip()
            continue

        # Parse the paragraph XML to find tracked changes
        para_xml = para._element
        parts = []
        has_changes = False

        for child in para_xml:
            tag = etree.QName(child.tag).localname if isinstance(child.tag, str) else ""

            if tag == "r":
                # Normal run
                text_els = child.findall(f"{{{WORD_NS}}}t")
                text = "".join(t.text or "" for t in text_els)
                if text:
                    parts.append(text)

            elif tag == "ins":
                # Insertion
                has_changes = True
                runs = child.findall(f".//{{{WORD_NS}}}t")
                text = "".join(t.text or "" for t in runs)
                if text:
                    parts.append(f"(+{text})")

            elif tag == "del":
                # Deletion
                has_changes = True
                del_runs = child.findall(f".//{{{WORD_NS}}}delText")
                text = "".join(t.text or "" for t in del_runs)
                if text:
                    parts.append(f"(~~{text}~~)")

        full_text = "".join(parts).strip()
        if not full_text:
            continue

        if has_changes:
            chunks.append(Chunk(
                text=full_text,
                section_heading=current_heading,
                is_redline=1,
                content_type="redline",
            ))
        else:
            chunks.append(Chunk(
                text=full_text,
                section_heading=current_heading,
                is_redline=1,
                content_type="text",
            ))

    return chunks


def process_docx_with_versions(
    path: str,
    doc_id: str,
) -> tuple[list[Chunk], list[Chunk]]:
    """
    Process DOCX into two parallel chunk sets.

    Returns:
        Tuple of (clean_chunks, redline_chunks).
        clean_chunks: Tracked changes accepted. is_redline=0.
        redline_chunks: Shows additions (+) and deletions (~~strikethrough~~) inline.
                        is_redline=1, redline_base_doc_id=doc_id of clean version.

    Args:
        path: Absolute path to the DOCX file.
        doc_id: Document identifier for metadata.

    Raises:
        FileNotFoundError: If path does not exist.
        docx.opc.exceptions.PackageNotFoundError: If file is not a valid DOCX.
    """
    file_path = Path(path)
    if not file_path.exists():
        raise FileNotFoundError(f"DOCX not found: {path}")

    logger.info(
        "Processing DOCX with versions",
        extra={"path": path, "doc_id": doc_id},
    )

    doc = docx.Document(path)

    # Clean version — tracked changes accepted
    clean_chunks = _extract_clean_text(doc)

    # Redline version — shows tracked changes inline
    redline_chunks = _extract_redline_text(doc)

    # Set redline_base_doc_id on all redline chunks
    for chunk in redline_chunks:
        chunk.redline_base_doc_id = doc_id

    logger.info(
        "DOCX processing complete",
        extra={
            "doc_id": doc_id,
            "clean_chunks": len(clean_chunks),
            "redline_chunks": len(redline_chunks),
        },
    )

    return clean_chunks, redline_chunks
